#!/usr/bin/env python3
"""Build the CCC v1 letter-library migration from the 38 master DOCX files.

The converter preserves every existing curly token and every fixed paragraph.
Only the highlighted ►► authoring/example blocks become explicit team-written
curlys, and operational instructions that are not part of a mailed letter are
removed. The generated migration carries hashes and token metadata so the Node
verification script can detect drift.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import uuid
from dataclasses import dataclass
from pathlib import Path

from docx import Document


NAMESPACE = uuid.UUID("df18cf41-a125-4a93-a52f-2b02d16df41c")
TOKEN_RE = re.compile(r"\{([^{}]+)\}")
FILE_RE = re.compile(r"^(ACC|COL|COMBO|CON|LP|DIRECT) - R(\d+) - (.+)\.docx$", re.I)
PREFIX_TO_FLOW = {
    "ACC": "accuracy",
    "COL": "collection",
    "COMBO": "combo",
    "CON": "consent",
    "LP": "late_pay",
    "DIRECT": "direct",
    "BONUS": "accuracy_solo",
}
EXPECTED_COUNTS = {
    "accuracy": 12,
    "collection": 10,
    "combo": 9,
    "consent": 3,
    "late_pay": 1,
    "direct": 2,
    "accuracy_solo": 1,
}
COMMON_HUMAN_TOKENS = {
    "DAMAGES": "damages",
    "LIST OF EXACT INACCURACIES": "personalization",
    "RANT SUMMARY": "personalization",
    "SPECIFIC INACCURATE CATEGORIES": "personalization",
    "OPTIONAL STRENGTHENER": "optional_strengthener",
    "PENALTY": "penalty",
    "CONSUMER STATEMENT": "consumer_statement",
}
DIRECT_HUMAN_TOKENS = {
    (1, "OPENING"): "damages",
    (1, "WHAT VERIFICATION ACTUALLY IS"): "personalization",
    (1, "WHAT IS NOT VERIFICATION"): "optional_strengthener",
    (1, "DEADLINE AND CLOSE"): "penalty",
    (2, "OPENING"): "personalization",
    (2, "THE LEGAL POSITION"): "optional_strengthener",
    (2, "DAMAGES"): "damages",
    (2, "THE CLOSE"): "penalty",
}
OPERATIONAL_PREFIXES = (
    "►► ATTACH ID + PROOF OF ADDRESS",
    "THIS LETTER GOES TO THE DEBT COLLECTOR",
)
EXAMPLE_ONLY_TOKENS = {"client_first_name", "bureau_name"}
NON_LEGAL_WORKFLOW_CORRECTIONS = {
    # CCC does not file a parallel CFPB complaint 7-14 days after each round.
    # This removes only the resulting factual assertion; the R11 statutes and
    # cease-communication theory remain exactly where the master places them.
    "Each was invoked in writing, each has a corresponding CFPB complaint on record, and none produced a compliant response.":
        "Each was invoked in writing, and none produced a compliant response.",
}


@dataclass(frozen=True)
class SourceTemplate:
    prefix: str
    round_number: int
    title: str
    path: Path

    @property
    def flow(self) -> str:
        return PREFIX_TO_FLOW[self.prefix]

    @property
    def key(self) -> str:
        label = "ACC-SOLO" if self.prefix == "BONUS" else self.prefix
        return f"{label}-R{self.round_number}-v1"

    @property
    def name(self) -> str:
        return self.path.stem


def discover_templates(source_dir: Path) -> list[SourceTemplate]:
    templates: list[SourceTemplate] = []
    for path in source_dir.glob("*.docx"):
        match = FILE_RE.match(path.name)
        if match:
            templates.append(SourceTemplate(
                match.group(1).upper(), int(match.group(2)), match.group(3), path,
            ))
        elif path.name.startswith("BONUS - "):
            templates.append(SourceTemplate(
                "BONUS", 1, path.stem.removeprefix("BONUS - "), path,
            ))
    templates.sort(key=lambda item: (item.flow, item.round_number, item.name))
    counts = {flow: sum(item.flow == flow for item in templates) for flow in EXPECTED_COUNTS}
    if len(templates) != 38 or counts != EXPECTED_COUNTS:
        raise RuntimeError(f"Expected the exact 38-template catalog; found {len(templates)} with {counts}")
    return templates


def instruction_token(template: SourceTemplate, paragraph: str) -> str | None:
    upper = paragraph.upper()
    if not upper.startswith("►►"):
        return None
    label = upper.removeprefix("►►").strip()
    label = label.removeprefix("WRITE THIS — ").strip()
    label = label.removeprefix("OPTIONAL STRENGTHENER — ").strip() if upper.startswith("►► OPTIONAL STRENGTHENER") else label
    if upper.startswith("►► OPTIONAL STRENGTHENER"):
        return "optional_strengthener"

    if template.flow == "direct":
        for (round_number, prefix), token in DIRECT_HUMAN_TOKENS.items():
            if template.round_number == round_number and label.startswith(prefix):
                return token
        return None

    for prefix, token in COMMON_HUMAN_TOKENS.items():
        if label.startswith(prefix):
            return token
    return None


def is_boundary(paragraph: str) -> bool:
    stripped = paragraph.strip()
    return stripped.startswith("►►") or stripped.startswith("— — —")


def extract_body(template: SourceTemplate) -> tuple[str, list[str], list[str], list[str]]:
    document = Document(template.path)
    if document.tables:
        raise RuntimeError(f"Tables are not supported in seeded letters: {template.path.name}")
    paragraphs = [paragraph.text.replace("\u00a0", " ").rstrip() for paragraph in document.paragraphs]
    source_text = "\n".join(paragraphs)
    source_tokens = sorted(set(TOKEN_RE.findall(source_text)))

    if template.flow == "accuracy_solo":
        marker = "— — — LETTER STARTS BELOW. DELETE EVERYTHING ABOVE THIS LINE BEFORE SENDING. — — —"
        try:
            paragraphs = paragraphs[paragraphs.index(marker) + 1:]
        except ValueError as exc:
            raise RuntimeError(f"Accuracy Solo letter-start marker missing in {template.path.name}") from exc

    output: list[str] = []
    human_tokens: list[str] = []
    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index].strip()
        if not paragraph:
            output.append("")
            index += 1
            continue

        if paragraph.startswith("►► PASTE SCREENSHOTS HERE"):
            output.append("{screenshots}")
            index += 1
            continue
        if paragraph.startswith(OPERATIONAL_PREFIXES):
            index += 1
            continue

        token = instruction_token(template, paragraph)
        if token:
            output.append(f"{{{token}}}")
            if token not in human_tokens:
                human_tokens.append(token)
            index += 1
            while index < len(paragraphs) and not is_boundary(paragraphs[index]):
                index += 1
            continue

        if paragraph.startswith("►►"):
            raise RuntimeError(f"Unmapped authoring instruction in {template.path.name}: {paragraph}")
        output.append(paragraphs[index])
        index += 1

    body = "\n".join(output)
    for source_text, replacement_text in NON_LEGAL_WORKFLOW_CORRECTIONS.items():
        body = body.replace(source_text, replacement_text)
    body = re.sub(r"\n{3,}", "\n\n", body).strip()
    body_tokens = sorted(set(TOKEN_RE.findall(body)))
    missing_source_tokens = sorted(set(source_tokens) - set(body_tokens))
    if set(missing_source_tokens) - EXAMPLE_ONLY_TOKENS:
        raise RuntimeError(f"Curly-token loss in {template.path.name}: {missing_source_tokens}")
    if "►►" in body or "EXAMPLE OF THE RIGHT LENGTH" in body:
        raise RuntimeError(f"Authoring/example text survived conversion in {template.path.name}")
    return body, source_tokens, human_tokens, missing_source_tokens


def sql_string(value: str) -> str:
    return "'" + value.replace("'", "''") + "'"


def migration_header() -> str:
    return """-- Install the 38 CCC-original v1 templates that model the course's
-- law/round structure without copying its wording. Apply after the template
-- table migration. Curly-token hashes are verified by the repository test.

alter table public.dispute_templates
  alter column created_by drop not null;

alter table public.dispute_templates
  drop constraint if exists dispute_templates_flow_code_check,
  drop constraint if exists dispute_templates_round_number_check;

alter table public.dispute_templates
  add constraint dispute_templates_flow_code_check
    check (flow_code in ('accuracy', 'collection', 'combo', 'consent', 'late_pay', 'direct', 'accuracy_solo')),
  add constraint dispute_templates_round_number_check
    check (
      (flow_code in ('accuracy', 'combo') and round_number between 1 and 12)
      or (flow_code = 'collection' and round_number between 1 and 10)
      or (flow_code = 'consent' and round_number between 1 and 3)
      or (flow_code in ('late_pay', 'direct') and round_number between 1 and 2)
      or (flow_code = 'accuracy_solo' and round_number = 1)
    );

alter table public.letters
  drop constraint if exists letters_dispute_flow_code_check;

alter table public.letters
  add constraint letters_dispute_flow_code_check
    check (dispute_flow_code is null or dispute_flow_code in (
      'accuracy', 'collection', 'combo', 'consent', 'late_pay', 'direct', 'accuracy_solo'
    ));

comment on column public.dispute_templates.created_by is
  'Null only for source-controlled CCC system templates; admin-created templates retain their auth user id.';

"""


def template_sql(template: SourceTemplate, ordinal: int) -> str:
    body, source_tokens, human_tokens, example_only_tokens = extract_body(template)
    tag = f"$ccc_{ordinal:02d}$"
    if tag in body:
        raise RuntimeError(f"Dollar quote collision in {template.path.name}")
    template_id = uuid.uuid5(NAMESPACE, template.key)
    body_hash = hashlib.sha256(body.encode("utf-8")).hexdigest()
    metadata = {
        "key": template.key,
        "flow": template.flow,
        "round": template.round_number,
        "source": template.path.name,
        "sourceTokens": source_tokens,
        "humanTokens": human_tokens,
        "exampleOnlyTokens": example_only_tokens,
        "bodySha256": body_hash,
    }
    notes = (
        f"CCC-original v1 rewrite modeled on the course's {template.prefix} R{template.round_number} "
        f"law and escalation purpose. Source: {template.path.name}. "
        f"Team fields: {', '.join(human_tokens) or 'none'}."
    )
    return f"""-- CCC-TEMPLATE {json.dumps(metadata, separators=(',', ':'), ensure_ascii=False)}
insert into public.dispute_templates (
  id, created_by, name, flow_code, round_number, bureau_code,
  version_label, body_text, notes, is_active
) values (
  '{template_id}', null, {sql_string(template.name)}, {sql_string(template.flow)},
  {template.round_number}, 'ALL', 'v1', {tag}{body}{tag}, {sql_string(notes)}, true
)
on conflict (id) do update set
  name = excluded.name,
  flow_code = excluded.flow_code,
  round_number = excluded.round_number,
  bureau_code = excluded.bureau_code,
  version_label = excluded.version_label,
  body_text = excluded.body_text,
  notes = excluded.notes,
  is_active = excluded.is_active,
  updated_at = now();

"""


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_dir", type=Path)
    parser.add_argument("output", type=Path)
    args = parser.parse_args()
    templates = discover_templates(args.source_dir)
    sql = migration_header() + "".join(
        template_sql(template, ordinal)
        for ordinal, template in enumerate(templates, start=1)
    )
    args.output.write_text(sql.rstrip() + "\n", encoding="utf-8")
    print(f"Wrote {len(templates)} CCC templates to {args.output}")


if __name__ == "__main__":
    main()
